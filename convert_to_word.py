import os
from docx import Document
from docx.shared import Pt
import re

def convert_md_to_docx(md_file, docx_file):
    doc = Document()
    
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
        
    lines = content.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            doc.add_paragraph('')
            continue
            
        # Headings
        if line.startswith('# '):
            doc.add_heading(line[2:], level=0)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=1)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=2)
            
        # Lists
        elif line.startswith('* ') or line.startswith('- '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif re.match(r'^\d+\.', line):
            doc.add_paragraph(line[re.search(r'\d+\.\s*', line).end():], style='List Number')
            
        # Normal text (parsing bold)
        else:
            p = doc.add_paragraph()
            # Simplistic bold parsing: **text**
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)

    # Styling tweak
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(11)

    doc.save(docx_file)

if __name__ == "__main__":
    convert_md_to_docx('GUIA_FUNDACION_TIENDA.md', 'GUIA_FUNDACION_TIENDA.docx')
